from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "專題_Ver.3.docx"


def compact(text: str) -> str:
    return " ".join(text.split())


def main() -> None:
    print("exists", DOCX.exists(), DOCX)
    doc = Document(str(DOCX))
    print("paragraphs", len(doc.paragraphs), "tables", len(doc.tables), "sections", len(doc.sections))
    print("inline_shapes", len(doc.inline_shapes))
    print("styles sample:")
    for style in list(doc.styles)[:100]:
        print(" ", style.name)
    print("\nFIRST_PARAS")
    shown = 0
    for i, para in enumerate(doc.paragraphs):
        txt = compact(para.text)
        if not txt:
            continue
        print(i, para.style.name, txt[:180])
        shown += 1
        if shown >= 80:
            break
    print("\nHEADINGS")
    for i, para in enumerate(doc.paragraphs):
        txt = compact(para.text)
        if not txt:
            continue
        style_name = para.style.name
        if (
            "Heading" in style_name
            or "標題" in style_name
            or txt.startswith(("第壹", "第一", "第貳", "第二", "第參", "第三", "第四", "第五", "第六"))
        ):
            print(i, style_name, txt[:220])


if __name__ == "__main__":
    main()
