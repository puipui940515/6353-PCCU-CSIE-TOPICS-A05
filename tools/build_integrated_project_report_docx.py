from __future__ import annotations

import json
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "專題_Ver.3.docx"
OUTPUT = ROOT / "專題_Ver.3_整合完善版_report.docx"
METRICS = ROOT / "report_metrics.json"
TB = ROOT / "report_tb_scalars.json"


def read_json(path: Path) -> dict:
    raw = path.read_bytes()
    for enc in ("utf-8", "utf-8-sig", "utf-16"):
        try:
            return json.loads(raw.decode(enc))
        except Exception:
            pass
    raise ValueError(f"Cannot decode {path}")


def pct(v: float) -> str:
    return f"{v * 100:.1f}%"


def anchor_element(anchor):
    return anchor._p if hasattr(anchor, "_p") else anchor._tbl


def insert_element_after(anchor, element):
    anchor_element(anchor).addnext(element)
    return element


def insert_paragraph_after(anchor, text: str = "", style: str | None = "Normal"):
    p = OxmlElement("w:p")
    anchor._p.addnext(p)
    para = anchor._parent.add_paragraph()
    para._p.getparent().remove(para._p)
    p.addprevious(para._p)
    para._p.getparent().remove(para._p)
    anchor._p.addnext(para._p)
    if style:
        try:
            para.style = style
        except KeyError:
            para.style = "Normal"
    if text:
        para.add_run(text)
    return para


def append_then_move_paragraph(doc: Document, anchor, text: str = "", style: str = "Normal"):
    para = doc.add_paragraph(text)
    try:
        para.style = style
    except KeyError:
        para.style = "Normal"
    para._p.getparent().remove(para._p)
    anchor_element(anchor).addnext(para._p)
    return para


def append_then_move_heading(doc: Document, anchor, text: str, level: int = 3):
    para = doc.add_heading(text, level=level)
    para._p.getparent().remove(para._p)
    anchor_element(anchor).addnext(para._p)
    return para


def format_body(para):
    para.paragraph_format.first_line_indent = Inches(0.32)
    para.paragraph_format.line_spacing = 1.5
    for run in para.runs:
        run.font.size = Pt(10.5)


def add_body(doc: Document, anchor, text: str):
    para = append_then_move_paragraph(doc, anchor, text, "Normal")
    format_body(para)
    return para


def add_figure_note(doc: Document, anchor, motivation: str, reason: str, conclusion: str):
    text = f"圖前說明：動機：{motivation} 理由：{reason} 結論：{conclusion}"
    para = append_then_move_paragraph(doc, anchor, text, "Normal")
    format_body(para)
    para.paragraph_format.keep_with_next = True
    return para


def add_caption(doc: Document, anchor, caption: str):
    para = append_then_move_paragraph(doc, anchor, caption, "Normal")
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in para.runs:
        run.font.size = Pt(10)
        run.bold = True
    return para


def add_picture(doc: Document, anchor, image: Path, width: float = 5.5):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if image.exists():
        para.add_run().add_picture(str(image), width=Inches(width))
    else:
        para.add_run(f"（缺圖：{image.relative_to(ROOT)}）")
    para._p.getparent().remove(para._p)
    anchor_element(anchor).addnext(para._p)
    return para


def apply_cell_borders(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "808080")


def add_table(doc: Document, anchor, headers: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1, cols=len(headers))
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    for row in table.rows:
        for cell in row.cells:
            apply_cell_borders(cell)
            for para in cell.paragraphs:
                para.style = "Normal"
                for run in para.runs:
                    run.font.size = Pt(9)
    table._tbl.getparent().remove(table._tbl)
    anchor_element(anchor).addnext(table._tbl)
    return table


def find_para(doc: Document, startswith: str):
    for p in doc.paragraphs:
        if p.text.strip().startswith(startswith):
            return p
    raise ValueError(f"Cannot find paragraph: {startswith}")


def replace_doc_text(doc: Document):
    replacements = {
        "GPU 算力實現零延遲的即時視覺推論與模擬": "GPU 算力實現低延遲的模型推論與模擬",
        "訓練集規模為 200,000 筆資料（隨機種子 0），驗證集為 20,000 筆（隨機種子 999）。": "訓練集規模為 200,000 筆資料（隨機種子 0），驗證集為 20,000 筆（隨機種子 999），資料內容已擴充至方位、距離、高度、不同 Hz、障礙物衰減係數與本體旋轉隨機化。",
        "感知網路採用多層感知機（MLP）架構": "感知網路採用多層感知機（MLP）架構，並以方位、距離與高度三個輸出 head 進行多任務學習",
    }
    for para in doc.paragraphs:
        updated = para.text
        for old, new in replacements.items():
            updated = updated.replace(old, new)
        if updated != para.text:
            para.text = updated


def add_block(doc: Document, anchor, items: list):
    current = anchor
    for item in items:
        kind = item[0]
        if kind == "heading":
            current = append_then_move_heading(doc, current, item[1], item[2])
        elif kind == "body":
            current = add_body(doc, current, item[1])
        elif kind == "table":
            current = add_table(doc, current, item[1], item[2])
        elif kind == "figure":
            _, motivation, reason, conclusion, image, caption, width = item
            note = add_figure_note(doc, current, motivation, reason, conclusion)
            pic = add_picture(doc, note, image, width)
            current = add_caption(doc, pic, caption)
    return current


def tb_last(tb: dict, run_contains: str, tag: str) -> str:
    matches = []
    for path, scalars in tb.get("perception", {}).items():
        if run_contains in path and tag in scalars:
            matches.append((scalars[tag]["last_step"], scalars[tag]["last_value"]))
    if not matches:
        return "未記錄"
    step, value = sorted(matches)[-1]
    return f"{value:.4f}（step {step:,}）"


def sac_rows(tb: dict):
    rows = []
    for path, scalars in sorted(tb.get("sac", {}).items()):
        reward = scalars.get("eval/mean_reward") if isinstance(scalars, dict) else None
        fps = scalars.get("time/fps") if isinstance(scalars, dict) else None
        if reward:
            parts = Path(path).parts
            run = parts[1] if len(parts) > 1 else path
            rows.append([run, f"{reward['last_value']:.2f}", f"{reward['last_step']:,}", f"{fps['last_value']:.0f}" if fps else "未記錄"])
    return rows[:4]


def main():
    shutil.copy2(SOURCE, OUTPUT)
    metrics = read_json(METRICS)
    tb = read_json(TB) if TB.exists() else {"perception": {}, "sac": {}}
    doc = Document(str(OUTPUT))
    replace_doc_text(doc)

    train = metrics["train_dataset"]
    evds = metrics["eval_dataset"]
    ev = metrics["eval"]

    acoustic_anchor = find_para(doc, "訓練資料批量輸入 GPU 運算")
    acoustic_items = [
        ("heading", "被動聲音資料生成與三維標籤補充", 3),
        ("body", "為使被動聲音感知模組能更接近真實機械臂工作情境，本研究將原本以方位為主的資料生成流程擴充為三維定位資料。每筆資料除了方位標籤外，亦包含距離分類、高度分類、聲源頻率、自身 yaw 旋轉與障礙物衰減係數。這些欄位可讓訓練不只學會「聲音從哪個方向來」，也能初步估計「距離多遠」與「高度落在哪個區間」。"),
        ("table", ["資料集", "樣本數", "特徵維度", "頻率範圍", "距離範圍", "高度範圍", "平均障礙物數"], [
            ["訓練集", f"{train['samples']:,}", str(train["feature_dim"]), f"{train['f0_hz']['min']:.0f}-{train['f0_hz']['max']:.0f} Hz", f"{train['source_ranges']['min']:.2f}-{train['source_ranges']['max']:.2f} m", f"{train['source_heights']['min']:.2f}-{train['source_heights']['max']:.2f} m", f"{train['obstacle_counts']['mean']:.2f}"],
            ["驗證集", f"{evds['samples']:,}", str(evds["feature_dim"]), f"{evds['f0_hz']['min']:.0f}-{evds['f0_hz']['max']:.0f} Hz", f"{evds['source_ranges']['min']:.2f}-{evds['source_ranges']['max']:.2f} m", f"{evds['source_heights']['min']:.2f}-{evds['source_heights']['max']:.2f} m", f"{evds['obstacle_counts']['mean']:.2f}"],
        ]),
        ("figure", "需要先確認麥克風陣列與聲源方位的相對關係，否則模型輸出的角度將難以對應到機械臂座標系。", "此圖以俯視角呈現麥克風陣列、聲源方向與方位角標籤的定義，可說明相位差特徵如何轉換為方位分類問題。", "資料生成的方位標籤與 demo 中的目標方向使用同一套幾何定義，因此訓練、檢查工具與 demo 顯示能維持一致。", ROOT / "detect/inspect_output/geometry.png", "圖 3-1 麥克風陣列與聲源方位幾何", 5.2),
        ("figure", "加入高度後，單純俯視圖不足以描述聲源與麥克風陣列之間的三維位置差。", "此圖用 3D 視角呈現聲源高度，對應新增的 height head，也能解釋 demo 中側面視角與高度變量的必要性。", "高度標籤已經進入資料生成與模型訓練流程，但高度正確率仍低於方位，後續需要更明確的垂直幾何線索。", ROOT / "detect/inspect_output/geometry_3d.png", "圖 3-2 三維聲源高度與麥克風陣列關係", 5.2),
        ("figure", "資料集若集中在少數角度，模型可能只記住常見方向而非真正學會定位。", "方位分布圖用來檢查訓練與驗證資料是否覆蓋完整 360 度，使模型能面對任意方向的聲源。", "目前資料分布覆蓋全方位，適合作為方位定位模型的基礎驗證資料。", ROOT / "detect/inspect_output/azimuth_distribution.png", "圖 3-3 驗證資料方位分布", 5.4),
        ("figure", "不同 Hz 的資料能降低模型對單一固定頻率的依賴，避免 demo 或真機環境頻率微小偏移時失效。", "頻率分布圖顯示資料生成在 38 kHz 至 42 kHz 間隨機抽樣，能讓模型學習更穩健的相位與能量關係。", "目前資料已具備頻率域隨機化，後續可再加入不同 SNR 與不同材質反射條件。", ROOT / "detect/inspect_output/f0_distribution.png", "圖 3-4 聲源頻率分布", 5.4),
        ("figure", "障礙物會造成不同麥克風通道收到的能量被不均勻削弱，是目前準確率下降的重要來源之一。", "障礙物衰減熱圖可檢查各通道隨機係數是否真的進入資料集，而不是只有紀錄 obstacle count。", "目前資料集已包含隨機障礙物係數，後續訓練可針對有障礙與無障礙樣本分別追蹤指標。", ROOT / "detect/inspect_output/obstacle_gain_heatmap.png", "圖 3-5 障礙物通道衰減係數熱圖", 5.4),
    ]
    add_block(doc, acoustic_anchor, acoustic_items)

    result_anchor = find_para(doc, "系統測試與實驗結果")
    result_items = [
        ("heading", "被動聲音感知模型實驗結果", 2),
        ("body", f"本節使用目前 checkpoint：{metrics['checkpoint']}，於 20,000 筆驗證資料上進行評估。模型採用 MLP backbone，並分別輸出方位、距離與高度分類結果。此模型屬於單窗推論架構，因此推論速度快，但尚不具備時間序列記憶力。若未來需要模型根據自身移動或旋轉的歷史變化修正判斷，應改為多幀輸入並加入 GRU、LSTM 或 Transformer Encoder。"),
        ("table", ["指標", "結果", "解讀"], [
            ["方位 10 度容差命中率", pct(ev["azimuth_hit_10deg"]), "用於判斷 demo 與控制端是否能取得可用方向"],
            ["方位平均誤差", f"{ev['azimuth_mean_err_deg']:.2f} 度", "越低代表方位估計越穩定"],
            ["距離分類正確率", pct(ev["range_acc"]), "4 類距離分類，隨機基準約 25%"],
            ["距離平均絕對誤差", f"{ev['range_mae_m']:.3f} m", "以分類 bin 中心估算距離誤差"],
            ["高度分類正確率", pct(ev["height_acc"]), "5 類高度分類，隨機基準約 20%"],
            ["高度平均絕對誤差", f"{ev['height_mae_m']:.3f} m", "以分類 bin 中心估算高度誤差"],
            ["單筆推論時間", f"{ev['latency_ms_per_sample']:.4f} ms", f"於 {ev['device']} 批次推論估算"],
            ["有障礙方位命中率", pct(ev["azimuth_hit_obstacle"]), "障礙物衰減下的方位定位表現"],
        ]),
        ("table", ["TensorBoard 指標", "主訓練 run 末值"], [
            ["train/loss", tb_last(tb, "detect\\runs\\gpu\\tb", "train/loss")],
            ["train/loss_azimuth", tb_last(tb, "detect\\runs\\gpu\\tb", "train/loss_azimuth")],
            ["train/loss_range", tb_last(tb, "detect\\runs\\gpu\\tb", "train/loss_range")],
            ["train/loss_height", tb_last(tb, "detect\\runs\\gpu\\tb", "train/loss_height")],
            ["eval/hit_rate", tb_last(tb, "detect\\runs\\gpu\\tb", "eval/hit_rate")],
            ["eval/range_hit_rate", tb_last(tb, "detect\\runs\\gpu\\tb", "eval/range_hit_rate")],
            ["eval/height_hit_rate", tb_last(tb, "detect\\runs\\gpu\\tb", "eval/height_hit_rate")],
        ]),
        ("figure", "方位分類模型即使有整體命中率，也需要檢查錯誤是否集中在特定角度。", "混淆矩陣能呈現真實方位與預測方位之間的對應關係，若出現斜線外的亮帶，代表存在系統性混淆。", "目前模型在部分角度仍有混淆，因此後續應針對盲區增加資料或調整麥克風幾何。", ROOT / "detect/inspect_model/confusion_azimuth.png", "圖 4-1 方位分類混淆矩陣", 5.4),
        ("figure", "只看 accuracy 無法知道模型內部特徵是否真的形成可分群結構。", "t-SNE 將高維 embedding 投影到平面，可檢查相近方位是否在特徵空間中保持鄰近。", "方位 embedding 已出現分群趨勢，表示模型有學到方向相關特徵，但群間重疊仍是誤差來源。", ROOT / "detect/inspect_model/embedding_tsne_azimuth.png", "圖 4-2 方位 embedding t-SNE 分布", 5.4),
        ("figure", "需要知道模型主要依賴哪些輸入特徵，才能判斷準確率不足是資料、幾何還是模型本身造成。", "saliency 圖顯示各相位差與能量比特徵對輸出影響的平均程度，可作為後續調整麥克風位置或特徵工程的依據。", "目前模型確實使用多個特徵來源，但高度 head 的可分性仍不足，顯示高度訊號需要更強的垂直資訊。", ROOT / "detect/inspect_model/saliency_mean.png", "圖 4-3 聲學特徵 saliency 分布", 5.4),
    ]
    add_block(doc, result_anchor, result_items)

    sac_anchor = find_para(doc, "SAC 強化學習訓練結果分析")
    rows = sac_rows(tb)
    sac_items = [
        ("body", "SAC 強化學習目前仍屬於課程式訓練的早期階段，主要用於驗證機械臂在 MuJoCo 環境中能否穩定互動、收集回饋並持續更新策略。由 TensorBoard 記錄可觀察到訓練流程能穩定執行，但 episode length 多維持在上限附近，代表完整 pick and place 任務仍需要更細緻的 reward shaping 與課程切換條件。"),
    ]
    if rows:
        sac_items.append(("table", ["SAC run", "最後平均 reward", "最後 step", "最後 FPS"], rows))
    add_block(doc, sac_anchor, sac_items)

    future_anchor = find_para(doc, "動態聲波導引抓取之進階驗證")
    future_items = [
        ("body", "目前感知模型為單窗 MLP，因此沒有明確的時間記憶能力；它能處理經過隨機化的自身 yaw 與位置變化，但無法像序列模型一樣利用「上一刻到下一刻」的連續變化修正估計。後續若要讓模型接受自身移動、旋轉造成的聲學變化，可在資料生成時輸出連續多幀特徵，並將網路改為 GRU、LSTM 或 Transformer Encoder，再把隱狀態或多幀 embedding 提供給 SAC 控制端。"),
    ]
    add_block(doc, future_anchor, future_items)

    doc.save(str(OUTPUT))
    print(OUTPUT)


if __name__ == "__main__":
    main()
