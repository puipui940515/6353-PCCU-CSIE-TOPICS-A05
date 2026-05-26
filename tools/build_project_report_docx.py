from __future__ import annotations

import json
import shutil
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "專題_Ver.3.docx"
OUTPUT = ROOT / "專題_Ver.3_含專案圖片與實驗數據_report.docx"
METRICS = ROOT / "report_metrics.json"
TB = ROOT / "report_tb_scalars.json"


def read_json(path: Path) -> dict:
    raw = path.read_bytes()
    for encoding in ("utf-8", "utf-8-sig", "utf-16"):
        try:
            return json.loads(raw.decode(encoding))
        except Exception:
            continue
    raise ValueError(f"Cannot decode JSON: {path}")


def pct(v: float) -> str:
    return f"{v * 100:.1f}%"


def num(v: float, digits: int = 3) -> str:
    return f"{v:.{digits}f}"


def add_intro(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.style = "Normal"
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10.5)
    p.paragraph_format.keep_with_next = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.style = "Normal"
    p.paragraph_format.first_line_indent = Inches(0.32)
    p.paragraph_format.line_spacing = 1.5


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    for style in ("表格", "Table Grid", "Normal Table"):
        try:
            table.style = style
            break
        except (KeyError, ValueError):
            continue
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            borders = tc_pr.first_child_found_in("w:tcBorders")
            if borders is None:
                borders = OxmlElement("w:tcBorders")
                tc_pr.append(borders)
            for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
                tag = f"w:{edge}"
                element = borders.find(qn(tag))
                if element is None:
                    element = OxmlElement(tag)
                    borders.append(element)
                element.set(qn("w:val"), "single")
                element.set(qn("w:sz"), "6")
                element.set(qn("w:space"), "0")
                element.set(qn("w:color"), "808080")
            for para in cell.paragraphs:
                para.style = "Normal"
                for run in para.runs:
                    run.font.size = Pt(9)


def add_picture_block(doc: Document, image: Path, intro: str, width_in: float = 5.7) -> None:
    if not image.exists():
        add_body(doc, f"（缺圖：{image.relative_to(ROOT)}）")
        return
    add_intro(doc, intro)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_together = True
    p.add_run().add_picture(str(image), width=Inches(width_in))


def tb_last(tb: dict, run_contains: str, tag: str) -> str:
    matches = []
    for path, scalars in tb.get("perception", {}).items():
        if run_contains in path and tag in scalars:
            matches.append((scalars[tag]["last_step"], scalars[tag]["last_value"]))
    if not matches:
        return "未記錄"
    step, value = sorted(matches)[-1]
    return f"{value:.4f}（step {step:,}）"


def sac_summary(tb: dict) -> list[list[str]]:
    rows = []
    for path, scalars in sorted(tb.get("sac", {}).items()):
        if not scalars:
            continue
        run = Path(path).parts[1] if len(Path(path).parts) > 1 else path
        reward = scalars.get("eval/mean_reward") or scalars.get("rollout/ep_rew_mean")
        fps = scalars.get("time/fps")
        if reward:
            rows.append([
                run,
                f"{reward['last_value']:.2f}",
                f"{int(reward['last_step']):,}",
                f"{fps['last_value']:.0f}" if fps else "未記錄",
            ])
    return rows[:4]


def replace_doc_text(doc: Document) -> list[str]:
    changes: list[str] = []
    replacements = {
        "GPU 算力實現零延遲的即時視覺推論與模擬": "GPU 算力實現低延遲的模型推論與模擬",
        "Gazebo Harmonic模擬器中建構物理環境": "MuJoCo 虛擬環境中建構物理環境，並保留 ROS 2/Gazebo 架構作為後續整合方向",
        "訓練集規模為 200,000 筆資料（隨機種子 0），驗證集為 20,000 筆（隨機種子 999）。": "訓練集規模為 200,000 筆資料（隨機種子 0），驗證集為 20,000 筆（隨機種子 999），資料內容已擴充至方位、距離、高度、不同 Hz、障礙物衰減係數與本體旋轉隨機化。",
        "感知網路採用多層感知機（MLP）架構": "感知網路採用多層感知機（MLP）架構，並以方位、距離與高度三個輸出 head 進行多任務學習",
    }
    for para in doc.paragraphs:
        original = para.text
        updated = original
        for old, new in replacements.items():
            updated = updated.replace(old, new)
        if updated != original:
            para.text = updated
            changes.append(updated)
    return changes


def main() -> None:
    shutil.copy2(SOURCE, OUTPUT)
    metrics = read_json(METRICS)
    tb = read_json(TB) if TB.exists() else {"perception": {}, "sac": {}}

    doc = Document(str(OUTPUT))
    changes = replace_doc_text(doc)

    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_heading("新增內容目錄", level=1)
    add_body(doc, "本節為依據目前專案程式、資料集、模型 checkpoint 與 TensorBoard 記錄所補入之圖表與實驗資料。原論文主體格式與既有章節架構保留，新增圖片皆採用圖說在上、圖片在下的論文格式。")
    toc_rows = [
        ["A.1", "專案圖表與資料集分布"],
        ["A.2", "被動聲音感知模型訓練與評估"],
        ["A.3", "強化學習運行效果與目前限制"],
        ["A.4", "版本修正與後續維護重點"],
    ]
    add_table(doc, ["章節", "內容"], toc_rows)

    doc.add_heading("專案圖表與資料集分布", level=1)
    add_body(doc, "本專案目前的被動聲音感知資料集由 6 顆麥克風的相位差與能量比形成 15 維特徵，並同時標記方位、距離與高度。為提升對實際環境的耐受度，資料生成已納入 38 kHz 至 42 kHz 的不同頻率、本體 yaw 旋轉，以及障礙物造成的通道衰減。")
    train = metrics["train_dataset"]
    eval_ds = metrics["eval_dataset"]
    add_table(
        doc,
        ["資料集", "樣本數", "特徵維度", "頻率範圍", "距離範圍", "高度範圍", "平均障礙物數"],
        [
            [
                "訓練集",
                f"{train['samples']:,}",
                str(train["feature_dim"]),
                f"{train['f0_hz']['min']:.0f}-{train['f0_hz']['max']:.0f} Hz",
                f"{train['source_ranges']['min']:.2f}-{train['source_ranges']['max']:.2f} m",
                f"{train['source_heights']['min']:.2f}-{train['source_heights']['max']:.2f} m",
                num(train["obstacle_counts"]["mean"], 2),
            ],
            [
                "驗證集",
                f"{eval_ds['samples']:,}",
                str(eval_ds["feature_dim"]),
                f"{eval_ds['f0_hz']['min']:.0f}-{eval_ds['f0_hz']['max']:.0f} Hz",
                f"{eval_ds['source_ranges']['min']:.2f}-{eval_ds['source_ranges']['max']:.2f} m",
                f"{eval_ds['source_heights']['min']:.2f}-{eval_ds['source_heights']['max']:.2f} m",
                num(eval_ds["obstacle_counts"]["mean"], 2),
            ],
        ],
    )

    images = [
        ("圖 A-1  麥克風陣列與聲源方位的俯視幾何，用於說明模型輸入特徵與方位標籤的來源。", ROOT / "detect/inspect_output/geometry.png", 5.4),
        ("圖 A-2  加入高度後的 3D 幾何關係，顯示聲源高度標籤與麥克風陣列的相對位置。", ROOT / "detect/inspect_output/geometry_3d.png", 5.4),
        ("圖 A-3  驗證資料的方位分布，用於確認資料生成沒有集中在少數角度。", ROOT / "detect/inspect_output/azimuth_distribution.png", 5.7),
        ("圖 A-4  驗證資料的距離分布，對應距離 head 的分類學習目標。", ROOT / "detect/inspect_output/range_distribution.png", 5.7),
        ("圖 A-5  驗證資料的高度分布，對應新增高度 head 的分類學習目標。", ROOT / "detect/inspect_output/height_distribution.png", 5.7),
        ("圖 A-6  聲源頻率分布，顯示資料集已涵蓋不同 Hz 以避免模型只記住固定頻率。", ROOT / "detect/inspect_output/f0_distribution.png", 5.7),
        ("圖 A-7  障礙物通道衰減係數熱圖，用於檢查障礙物隨機化是否實際影響各麥克風通道。", ROOT / "detect/inspect_output/obstacle_gain_heatmap.png", 5.7),
    ]
    for intro, image, width in images:
        add_picture_block(doc, image, intro, width)

    doc.add_heading("被動聲音感知模型訓練與評估", level=1)
    ev = metrics["eval"]
    add_body(doc, "目前使用的感知模型 checkpoint 為 "
             f"{metrics['checkpoint']}。模型為無時間記憶的 stateless MLP，具有方位、距離與高度三個輸出 head；若要讓模型具有明確記憶力，後續應改為以連續多幀特徵輸入 GRU、LSTM 或 Transformer Encoder，並在訓練資料中保留時間序列。")
    add_table(
        doc,
        ["項目", "結果", "說明"],
        [
            ["方位 10 度容差命中率", pct(ev["azimuth_hit_10deg"]), "評估方位定位是否可提供 demo 與控制端使用"],
            ["方位 exact class accuracy", pct(ev["azimuth_acc"]), "72 類、每類 5 度的嚴格分類正確率"],
            ["方位平均誤差", f"{ev['azimuth_mean_err_deg']:.2f} 度", "越低代表定位穩定性越高"],
            ["距離分類正確率", pct(ev["range_acc"]), "4 類距離 head，隨機基準約 25%"],
            ["距離平均絕對誤差", f"{ev['range_mae_m']:.3f} m", "以各 bin 中心估計距離誤差"],
            ["高度分類正確率", pct(ev["height_acc"]), "5 類高度 head，隨機基準約 20%"],
            ["高度平均絕對誤差", f"{ev['height_mae_m']:.3f} m", "以各 bin 中心估計高度誤差"],
            ["單筆推論時間", f"{ev['latency_ms_per_sample']:.4f} ms", f"於 {ev['device']} 批次推論估算"],
            ["無障礙方位命中率", pct(ev["azimuth_hit_clear"]), "驗證障礙物前的基準表現"],
            ["有障礙方位命中率", pct(ev["azimuth_hit_obstacle"]), "障礙物衰減下仍保留可用定位能力"],
        ],
    )
    add_table(
        doc,
        ["TensorBoard 指標", "目前主訓練 run 末值"],
        [
            ["train/loss", tb_last(tb, "detect\\runs\\gpu\\tb", "train/loss")],
            ["train/loss_azimuth", tb_last(tb, "detect\\runs\\gpu\\tb", "train/loss_azimuth")],
            ["train/loss_range", tb_last(tb, "detect\\runs\\gpu\\tb", "train/loss_range")],
            ["train/loss_height", tb_last(tb, "detect\\runs\\gpu\\tb", "train/loss_height")],
            ["eval/hit_rate", tb_last(tb, "detect\\runs\\gpu\\tb", "eval/hit_rate")],
            ["eval/range_hit_rate", tb_last(tb, "detect\\runs\\gpu\\tb", "eval/range_hit_rate")],
            ["eval/height_hit_rate", tb_last(tb, "detect\\runs\\gpu\\tb", "eval/height_hit_rate")],
        ],
    )
    model_images = [
        ("圖 A-8  方位分類混淆矩陣，用於觀察模型是否出現固定角度偏誤或前後混淆。", ROOT / "detect/inspect_model/confusion_azimuth.png", 5.7),
        ("圖 A-9  t-SNE 方位嵌入圖，用於檢查模型內部特徵是否依方位形成可分群結構。", ROOT / "detect/inspect_model/embedding_tsne_azimuth.png", 5.7),
        ("圖 A-10  平均 saliency 分布，用於觀察哪些聲學特徵對模型判斷影響較高。", ROOT / "detect/inspect_model/saliency_mean.png", 5.7),
        ("圖 A-11  訓練與驗證樣本的最近鄰距離分布，用於檢查驗證集是否過度貼近訓練集。", ROOT / "detect/inspect_model/nearest_neighbor_dist.png", 5.7),
        ("圖 A-12  第一層權重視覺化，用於檢查模型是否有效利用相位差與能量比特徵。", ROOT / "detect/inspect_model/weight_layer1.png", 5.7),
    ]
    for intro, image, width in model_images:
        add_picture_block(doc, image, intro, width)

    doc.add_heading("強化學習運行效果與目前限制", level=1)
    add_body(doc, "SAC 訓練仍處於 Stage 1 課程式訓練階段，TensorBoard 記錄顯示訓練具備可持續運行能力，但平均 episode length 多維持在 500，表示完整 pick and place 仍需要後續 reward shaping、課程階段切換與成功條件設計。")
    rows = sac_summary(tb)
    if rows:
        add_table(doc, ["SAC run", "最後平均 reward", "最後 step", "最後 FPS"], rows)
    else:
        add_body(doc, "本次未讀取到 SAC TensorBoard scalar，因此僅保留模型檔與訓練設定作為實驗附件。")

    doc.add_heading("版本修正與後續維護重點", level=1)
    add_body(doc, "本次文件同步了目前專案狀態：感知資料生成加入高度、不同 Hz、自身旋轉與障礙物衰減；demo 已加入距離顯示與 T 鍵側視角切換；inspect 與 demo 的可見輸出改為 ASCII 英文以避免終端亂碼。")
    if changes:
        add_table(doc, ["本次微調文字"], [[c[:220]] for c in changes[:6]])
    add_table(
        doc,
        ["問題", "目前處理方式", "後續建議"],
        [
            ["高度正確率仍偏低", "先以分類 head 建立可用基準", "增加垂直麥克風間距、提高高度樣本密度或改為序列模型"],
            ["障礙物會降低方位命中", "資料生成已加入隨機衰減係數", "加入幾何遮蔽模型並分別記錄 obstacle/no-obstacle 指標"],
            ["模型尚無記憶力", "目前 MLP 單窗推論、延遲低", "改用多幀資料與 GRU/LSTM/Transformer Encoder"],
            ["SAC 任務尚未完整收斂", "Stage 1 可持續訓練", "調整 reward shaping、成功獎勵與 curriculum 門檻"],
        ],
    )

    doc.save(str(OUTPUT))
    print(OUTPUT)


if __name__ == "__main__":
    main()
